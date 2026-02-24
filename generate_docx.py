"""
Generate UrbanFlow Technical Documentation DOCX
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

def set_cell_shading(cell, color_hex):
    """Set background color of a table cell."""
    shading = cell._element.get_or_add_tcPr()
    shade = shading.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear',
        qn('w:color'): 'auto',
        qn('w:fill'): color_hex
    })
    shading.append(shade)

def add_styled_heading(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        if level == 1:
            run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)
            run.font.size = Pt(22)
        elif level == 2:
            run.font.color.rgb = RGBColor(0x2C, 0x5F, 0x8A)
            run.font.size = Pt(16)
        elif level == 3:
            run.font.color.rgb = RGBColor(0x34, 0x7A, 0xB5)
            run.font.size = Pt(13)
    return heading

def add_para(doc, text, bold=False, italic=False, size=11, color=None, alignment=None, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = 'Calibri'
    if color:
        run.font.color.rgb = RGBColor(*color)
    if alignment:
        p.alignment = alignment
    p.paragraph_format.space_after = Pt(space_after)
    return p

def add_bullet(doc, text, bold_prefix=None, level=0):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.size = Pt(11)
        run = p.add_run(text)
        run.font.size = Pt(11)
    else:
        run = p.add_run(text)
        run.font.size = Pt(11)
    return p

def create_table(doc, headers, rows, col_widths=None, header_color='1E3A5F'):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, header_color)

    # Data rows
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r + 1].cells[c]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(10)
            if r % 2 == 0:
                set_cell_shading(cell, 'F0F4F8')

    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)

    doc.add_paragraph()
    return table


def generate_document():
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # ========== TITLE PAGE ==========
    for _ in range(4):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('UrbanFlow')
    run.font.size = Pt(42)
    run.bold = True
    run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('AI-Powered Adaptive Traffic Signal Control System')
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x2C, 0x5F, 0x8A)

    doc.add_paragraph()

    desc = doc.add_paragraph()
    desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = desc.add_run('Technical Documentation\nHow UrbanFlow Solves Traffic Problems')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    for _ in range(6):
        doc.add_paragraph()

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run('Hackathon 2026')
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.add_page_break()

    # ========== TABLE OF CONTENTS ==========
    add_styled_heading(doc, 'Table of Contents', 1)
    toc_items = [
        '1. Executive Summary',
        '2. The Traffic Problem',
        '3. How UrbanFlow Solves It — The 6-Phase Pipeline',
        '   3.1 Phase 1: Data Collection (The Eyes)',
        '   3.2 Phase 2: Perception (Computer Vision)',
        '   3.3 Phase 3: Lane Assignment & Metrics',
        '   3.4 Phase 4: AI Decision Making (The Brain)',
        '   3.5 Phase 5: Signal Control (The Action)',
        '   3.6 Phase 6: Dashboard & Analytics (The Display)',
        '4. Complete File Architecture & Mapping',
        '   4.1 Backend Files',
        '   4.2 Frontend Files',
        '   4.3 Configuration & Infrastructure Files',
        '5. Technical Skills & Technologies Used',
        '6. APIs — Internal & External',
        '   6.1 REST APIs (FastAPI)',
        '   6.2 WebSocket APIs (Socket.IO)',
        '   6.3 Mobile Camera Server APIs (Flask)',
        '   6.4 External Libraries & APIs',
        '7. Data Flow — End-to-End Cycle',
        '8. Three Operating Modes',
        '9. Key Innovation — Traditional vs. UrbanFlow',
        '10. Conclusion',
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(2)
        for run in p.runs:
            run.font.size = Pt(11)

    doc.add_page_break()

    # ========== 1. EXECUTIVE SUMMARY ==========
    add_styled_heading(doc, '1. Executive Summary', 1)
    add_para(doc,
        'UrbanFlow is a smart, AI-powered traffic management system designed to replace traditional fixed-timer '
        'traffic signals with an intelligent, adaptive solution. It uses computer vision (via mobile phone cameras) '
        'to see real-time traffic, a reinforcement learning AI agent to decide optimal signal timings, and a real-time '
        'dashboard to visualize everything. The result is an estimated 30–40% reduction in waiting time, fuel waste, '
        'and emissions at managed intersections.')

    add_para(doc, 'The system operates through a seamless 6-phase pipeline:', bold=True)
    phases = [
        ('SEES ', 'traffic using a mobile phone camera (Computer Vision)'),
        ('THINKS ', 'about the best signal timing using AI (Reinforcement Learning)'),
        ('ACTS ', 'by optimizing traffic signal phases to reduce congestion'),
        ('SHOWS ', 'everything on a real-time dashboard'),
    ]
    for prefix, text in phases:
        add_bullet(doc, text, bold_prefix=prefix)

    doc.add_page_break()

    # ========== 2. THE TRAFFIC PROBLEM ==========
    add_styled_heading(doc, '2. The Traffic Problem', 1)
    add_para(doc,
        'Traditional traffic signal systems suffer from fundamental limitations that lead to urban congestion, '
        'wasted time, and environmental damage:')

    problems = [
        ('Fixed-Timer Signals: ', 'Conventional traffic lights operate on pre-set timers (e.g., 30 seconds green, 30 seconds red) regardless of actual traffic conditions. This means the same timing is used at 3 AM (empty roads) as at 5 PM (rush hour).'),
        ('No Traffic Awareness: ', 'Traditional systems are completely blind to the actual number of vehicles on the road. They cannot adapt to real-time demand.'),
        ('Unnecessary Waiting: ', 'Drivers wait at red lights even when no vehicles are crossing from the other direction, leading to frustration and wasted time.'),
        ('Long Queues on Busy Lanes: ', 'When one direction has heavy traffic and another is empty, both get equal green time, causing queues to build on the busy lanes.'),
        ('Fuel Waste & Emissions: ', 'Idling vehicles at unnecessarily long red lights consume fuel and emit CO₂, contributing to air pollution and climate change.'),
        ('No Data or Insights: ', 'Traffic authorities have no visibility into real-time conditions, making it impossible to optimize or respond proactively.'),
    ]
    for prefix, text in problems:
        add_bullet(doc, text, bold_prefix=prefix)

    doc.add_page_break()

    # ========== 3. HOW URBANFLOW SOLVES IT ==========
    add_styled_heading(doc, '3. How UrbanFlow Solves It — The 6-Phase Pipeline', 1)

    # --- Phase 1 ---
    add_styled_heading(doc, '3.1 Phase 1: Data Collection (The Eyes)', 2)
    add_para(doc,
        'UrbanFlow turns any smartphone into a traffic sensor. A user opens the camera server URL in their phone\'s '
        'browser, which accesses the rear camera and streams video frames in real-time to the backend server.')

    add_para(doc, 'How it works:', bold=True)
    steps_p1 = [
        'The phone opens http://<PC-IP>:5000 in the browser.',
        'The browser requests camera access and starts capturing video.',
        'Every ~200ms, JavaScript on the phone captures a video frame, converts it to Base64 JPEG, and sends it via Socket.IO to the Camera Server (Flask, Port 5000).',
        'The Camera Server decodes the Base64 data into an OpenCV numpy image array for processing.',
    ]
    for s in steps_p1:
        add_bullet(doc, s)

    add_para(doc, 'Files Involved:', bold=True)
    create_table(doc,
        ['File', 'Purpose'],
        [
            ['backend/app/services/mobile_camera/mobile_camera_server.py', 'Flask + Socket.IO server that serves the camera HTML page to the phone, receives frames via WebSocket, runs vehicle/signal detection, and exposes REST API endpoints'],
            ['backend/app/services/vision/camera_stream.py', 'CameraStream & MobileCamera classes — manages video input from phones/webcams/files, handles Base64 decoding, frame queuing, FPS calculation, and recording'],
        ],
        col_widths=[7, 11]
    )

    # --- Phase 2 ---
    add_styled_heading(doc, '3.2 Phase 2: Perception (Computer Vision)', 2)
    add_para(doc,
        'Each incoming frame is processed through two parallel detectors that extract structured traffic data from raw video:')

    add_styled_heading(doc, 'A. Vehicle Detector', 3)
    add_para(doc,
        'Detects and classifies vehicles in each frame using either YOLOv8 deep learning (when available) or '
        'MOG2 Background Subtraction (as fallback). Outputs per-vehicle bounding boxes, classifications (car, '
        'motorcycle, bus, truck), lane assignments, speed estimates, and queue lengths.')
    add_para(doc, 'Detection Pipeline:', bold=True)
    vd_steps = [
        'Background subtraction (MOG2) learns "empty road" and detects new/moving pixels.',
        'Morphological cleanup removes noise from the binary mask.',
        'Contour detection finds vehicle-shaped blobs (area > 500px², aspect ratio 0.2–4.0).',
        'Each detected vehicle gets a bounding box, center point, and classification.',
        'Frame-to-frame tracking estimates vehicle speeds.',
    ]
    for s in vd_steps:
        add_bullet(doc, s)

    add_styled_heading(doc, 'B. Signal Detector', 3)
    add_para(doc,
        'Detects traffic light state (Red, Yellow, Green) using HSV color space masking and Hough Circle Transform. '
        'Tracks phase durations and signal cycle times for the RL agent.')
    add_para(doc, 'Detection Pipeline:', bold=True)
    sd_steps = [
        'Convert frame from BGR to HSV color space.',
        'Apply separate color masks for Red (H: 0–10 & 160–180), Yellow, and Green.',
        'Detect circular shapes in each mask using HoughCircles.',
        'Determine signal state based on strongest detection with confidence scoring.',
        'Track phase transitions and calculate cycle timings.',
    ]
    for s in sd_steps:
        add_bullet(doc, s)

    add_para(doc, 'Files Involved:', bold=True)
    create_table(doc,
        ['File', 'Purpose'],
        [
            ['backend/app/services/vision/vehicle_detector.py', 'VehicleDetector class — YOLOv8 / MOG2 vehicle detection, lane assignment, speed estimation, queue length calculation, vehicle type counting, detection history export, and RL state conversion'],
            ['backend/app/services/vision/signal_detector.py', 'SignalDetector class — HSV color masking, Hough Circle detection, signal state classification, phase duration tracking, cycle time calculation, and color calibration'],
        ],
        col_widths=[7, 11]
    )

    # --- Phase 3 ---
    add_styled_heading(doc, '3.3 Phase 3: Lane Assignment & Metrics', 2)
    add_para(doc,
        'After detection, vehicles are assigned to 4 directional lanes (North, South, East, West) based on their '
        'position in the camera frame. The frame is divided into quadrants:')
    lane_rules = [
        'center_x < width/2 AND center_y < height/2 → NORTH',
        'center_x >= width/2 AND center_y < height/2 → EAST',
        'center_x < width/2 AND center_y >= height/2 → WEST',
        'center_x >= width/2 AND center_y >= height/2 → SOUTH',
    ]
    for r in lane_rules:
        add_bullet(doc, r)

    add_para(doc, 'Per-Lane Metrics Calculated:', bold=True)
    metrics_list = [
        ('Vehicle Count: ', 'Total vehicles detected in that lane'),
        ('Queue Length: ', 'Number of stopped/slow-moving vehicles (speed below threshold)'),
        ('Average Speed: ', 'Mean speed in m/s, estimated via frame-to-frame tracking'),
    ]
    for prefix, text in metrics_list:
        add_bullet(doc, text, bold_prefix=prefix)

    add_para(doc, 'The DataProcessor then converts all this into a unified state and computes derived metrics:', bold=False)
    derived = [
        ('Congestion Index: ', 'Normalized 0–1 score based on total queue length vs. maximum capacity'),
        ('Estimated Throughput: ', 'Total vehicles × normalized average speed'),
        ('Temporal Features: ', 'Frame-to-frame changes in queue and vehicle counts for trend awareness'),
    ]
    for prefix, text in derived:
        add_bullet(doc, text, bold_prefix=prefix)

    add_para(doc, 'Files Involved:', bold=True)
    create_table(doc,
        ['File', 'Purpose'],
        [
            ['backend/app/services/vision/data_processor.py', 'DataProcessor class — bridges vision output to RL input. Processes signal + vehicle data into unified state, normalizes features into a 32-dimensional RL state vector, calculates rewards, manages state/signal/vehicle history, and exports training data'],
        ],
        col_widths=[7, 11]
    )

    # --- Phase 4 ---
    add_styled_heading(doc, '3.4 Phase 4: AI Decision Making (The Brain)', 2)
    add_para(doc,
        'The core intelligence of UrbanFlow is a Reinforcement Learning (RL) agent that learns optimal traffic signal '
        'control policies through experience. It uses an Actor-Critic (A2C) architecture implemented in PyTorch.')

    add_para(doc, 'Neural Network Architecture:', bold=True)
    nn_rows = [
        ['Input Layer', '32 neurons', 'Receives the normalized state vector (queue, vehicles, speed, phase × 4 lanes + global metrics + temporal features)'],
        ['Hidden Layer', '128 neurons + ReLU', 'Learns complex traffic patterns and correlations'],
        ['Actor Head', '4 outputs + Softmax', 'Outputs probability distribution over 4 signal phases (Green, Yellow, Red, Extended Green)'],
        ['Critic Head', '1 output (Value)', 'Estimates how "good" the current traffic state is for learning'],
    ]
    create_table(doc,
        ['Layer', 'Size', 'Description'],
        nn_rows,
        col_widths=[3, 4, 11]
    )

    add_para(doc, 'Reward Function:', bold=True)
    add_para(doc, '    reward = -(waiting_time + 10 × stopped_vehicles)', italic=True, size=11)
    add_para(doc, 'The AI learns to maximize this reward, effectively minimizing waiting time and the number of stopped vehicles. Training runs for 100+ episodes with 1000 steps per episode.')

    add_para(doc, 'Additionally, a TrafficPredictor class uses Facebook Prophet for time-series forecasting of future traffic volumes.', italic=True)

    add_para(doc, 'Files Involved:', bold=True)
    create_table(doc,
        ['File', 'Purpose'],
        [
            ['backend/app/services/ai_engine/rl_agent.py', 'RLAgent + ActorCritic classes — PyTorch A2C neural network, action selection via categorical sampling, experience storage (log_probs, values, rewards, masks), discounted reward computation, advantage-based policy gradient updates, model save/load'],
            ['backend/app/services/ai_engine/predictor.py', 'TrafficPredictor class — Facebook Prophet time-series model, trains on historical vehicle counts, predicts future traffic with confidence intervals'],
        ],
        col_widths=[7, 11]
    )

    # --- Phase 5 ---
    add_styled_heading(doc, '3.5 Phase 5: Signal Control (The Action)', 2)
    add_para(doc,
        'The AI\'s chosen action is translated into a signal phase change. UrbanFlow supports three execution modes:')

    control_modes = [
        ('SUMO Mode: ', 'Uses the TraCI API to change actual traffic signal phases in the SUMO microscopic traffic simulator (traci.trafficlight.setPhase).'),
        ('Mock Mode: ', 'Stores the phase in memory and broadcasts to the dashboard — used for testing without SUMO installed.'),
        ('Real-World Mode (Future): ', 'Could send commands to IoT traffic controllers via MQTT, HTTP, or serial port.'),
    ]
    for prefix, text in control_modes:
        add_bullet(doc, text, bold_prefix=prefix)

    add_para(doc, 'Signal Phases: Green (0) → Yellow (1) → Red (2) → Extended Green (3)', bold=True)

    add_para(doc, 'Files Involved:', bold=True)
    create_table(doc,
        ['File', 'Purpose'],
        [
            ['backend/app/services/sumo/environment.py', 'SumoEnvironment class — interfaces with SUMO via TraCI API. Starts/stops simulation, advances steps, reads intersection state (queue length, vehicle count, avg speed, phase), applies signal actions, calculates rewards. Falls back to mock mode if SUMO not installed'],
            ['backend/app/services/vision/hybrid_environment.py', 'HybridEnvironment class — combines SUMO simulation + real-world camera data in configurable ratios (simulation/real_world/hybrid modes), processes states from both sources, supports weighted blending (70% sim + 30% camera)'],
            ['backend/app/api/routes/simulation.py', 'SimulationManager + API routes — manages the simulation loop, integrates HybridEnvironment + RLAgent, broadcasts state via WebSocket every 2 seconds. Endpoints: start, stop, status, camera frame ingestion, test broadcast'],
        ],
        col_widths=[7, 11]
    )

    # --- Phase 6 ---
    add_styled_heading(doc, '3.6 Phase 6: Dashboard & Analytics (The Display)', 2)
    add_para(doc,
        'The React.js frontend provides a real-time, interactive dashboard for traffic operators to monitor, '
        'understand, and control the entire system.')

    add_para(doc, 'Dashboard Components:', bold=True)
    dashboard_components = [
        ('TrafficMap: ', 'Interactive Leaflet map showing intersection locations with color-coded markers (green/yellow/red based on signal phase). Uses user geolocation to generate nearby intersections. Custom traffic light SVG icons.'),
        ('TrafficCharts: ', 'Real-time Recharts visualizations — bar chart for per-intersection vehicle throughput (color-coded by congestion level) and area chart for queue length trends over time.'),
        ('ExplainabilityPanel: ', 'XAI component showing WHY the AI made its decision — displays the action name, confidence score, natural language reasoning (with typewriter animation), and SHAP-like feature importance bars.'),
        ('CameraFeed: ', 'Live traffic camera feed embedded via iframe from the Mobile Camera Server (port 5000). Shows connection status, source info, and links to the full detection dashboard.'),
    ]
    for prefix, text in dashboard_components:
        add_bullet(doc, text, bold_prefix=prefix)

    add_para(doc, 'Files Involved:', bold=True)
    create_table(doc,
        ['File', 'Purpose'],
        [
            ['frontend/src/pages/Dashboard.jsx', 'Main dashboard page — assembles TrafficMap, TrafficCharts, ExplainabilityPanel, CameraFeed. Shows stat cards (total vehicles, avg wait time, active intersections, AI decisions). Start/stop simulation controls.'],
            ['frontend/src/pages/Analytics.jsx', 'Dedicated analytics page — KPI cards (avg wait time reduction, throughput, emissions saved), wait time optimization area chart, emissions savings bar chart'],
            ['frontend/src/pages/Home.jsx', 'Landing/home page with navigation'],
            ['frontend/src/components/Dashboard/TrafficMap.jsx', 'Leaflet map with traffic light markers, user location, congestion circles, click-to-view popups'],
            ['frontend/src/components/Analytics/TrafficCharts.jsx', 'Recharts bar chart (throughput) + area chart (queue trends) with real-time data from TrafficContext'],
            ['frontend/src/components/Dashboard/ExplainabilityPanel.jsx', 'AI explainability with typewriter effect, confidence bars, SHAP feature importance visualization (Framer Motion animations)'],
            ['frontend/src/components/Dashboard/CameraFeed.jsx', 'Live camera iframe, connection status indicators, link to camera dashboard'],
            ['frontend/src/context/TrafficContext.jsx', 'React Context provider — manages global state for intersections, historical data, AI decisions, emergency alerts, metrics. Consumes WebSocket events and falls back to mock data when disconnected'],
            ['frontend/src/services/websocket.js', 'WebSocketService class — Socket.IO client connecting to backend (port 8000). Handles traffic_update, ai_decision, prediction_update, emergency_alert, and metrics events with auto-reconnection'],
            ['frontend/src/services/api.js', 'Axios HTTP client — REST API calls for getAllIntersections, getIntersectionStatus, updateSignal, simulation start/stop/status'],
            ['frontend/src/App.jsx', 'Root component — React Router (Home, Dashboard, Analytics), MUI dark theme with Inter font, CssBaseline'],
            ['frontend/src/index.js', 'App entry point — BrowserRouter, React 18 root render'],
        ],
        col_widths=[7, 11]
    )

    doc.add_page_break()

    # ========== 4. FILE ARCHITECTURE ==========
    add_styled_heading(doc, '4. Complete File Architecture & Mapping', 1)

    add_styled_heading(doc, '4.1 Backend Files', 2)
    create_table(doc,
        ['File Path', 'Key Classes/Functions', 'Problem Solved'],
        [
            ['app/main.py', 'FastAPI app, Socket.IO ASGI wrapper, connect/disconnect events, broadcast_loop startup', 'Central server orchestration — CORS, routing, WebSocket lifecycle'],
            ['app/core/socket_manager.py', 'AsyncServer (sio), broadcast_traffic_update(), broadcast_ai_decision()', 'Real-time push communication to all connected dashboard clients'],
            ['app/api/routes/traffic.py', 'get_all_intersections(), get_intersection_status(), update_signal()', 'REST API for reading/writing intersection state and signal control'],
            ['app/api/routes/simulation.py', 'SimulationManager, HybridAdapter, broadcast_loop(), start/stop/status/camera endpoints', 'Simulation lifecycle, RL agent integration, 2-second broadcast loop'],
            ['app/models/traffic_models.py', 'TrafficData, SignalState, IntersectionStatus (Pydantic)', 'Data validation and serialization for API request/response'],
            ['app/services/vision/vehicle_detector.py', 'VehicleDetector (YOLOv8/MOG2)', 'Detecting and tracking vehicles from camera frames'],
            ['app/services/vision/signal_detector.py', 'SignalDetector (HSV + Hough)', 'Detecting traffic signal state from camera frames'],
            ['app/services/vision/camera_stream.py', 'CameraStream, MobileCamera', 'Managing video input from phones/webcams/files'],
            ['app/services/vision/data_processor.py', 'DataProcessor (32-dim state vector)', 'Bridging vision output to RL agent input'],
            ['app/services/vision/hybrid_environment.py', 'HybridEnvironment (sim + real)', 'Combining simulation and camera data'],
            ['app/services/sumo/environment.py', 'SumoEnvironment (TraCI)', 'SUMO traffic simulation interface'],
            ['app/services/ai_engine/rl_agent.py', 'RLAgent, ActorCritic (A2C)', 'AI-based adaptive signal timing optimization'],
            ['app/services/ai_engine/predictor.py', 'TrafficPredictor (Prophet)', 'Forecasting future traffic volumes'],
            ['app/services/mobile_camera/mobile_camera_server.py', 'Flask + SocketIO server, camera HTML page', 'Phone-to-server video streaming and detection'],
        ],
        col_widths=[5, 5, 7]
    )

    add_styled_heading(doc, '4.2 Frontend Files', 2)
    create_table(doc,
        ['File Path', 'Key Components', 'Problem Solved'],
        [
            ['src/App.jsx', 'Routes, MUI ThemeProvider', 'App routing and dark theme setup'],
            ['src/index.js', 'React 18 root, BrowserRouter', 'App entry point and routing provider'],
            ['src/pages/Dashboard.jsx', 'DashboardContent, StatCard', 'Main control center for traffic monitoring'],
            ['src/pages/Analytics.jsx', 'KPICard, ChartCard, Analytics', 'Historical trends and environmental impact'],
            ['src/pages/Home.jsx', 'Home', 'Landing page and navigation'],
            ['src/components/Dashboard/TrafficMap.jsx', 'TrafficMap, MapController', 'Geospatial visualization of intersections'],
            ['src/components/Analytics/TrafficCharts.jsx', 'TrafficCharts', 'Real-time vehicle throughput and queue trends'],
            ['src/components/Dashboard/ExplainabilityPanel.jsx', 'ExplainabilityPanel, TypewriterEffect', 'AI transparency and decision explanation'],
            ['src/components/Dashboard/CameraFeed.jsx', 'CameraFeed', 'Live camera monitoring from dashboard'],
            ['src/context/TrafficContext.jsx', 'TrafficProvider, useTraffic', 'Global state management with WebSocket sync'],
            ['src/services/websocket.js', 'WebSocketService', 'Real-time bidirectional communication'],
            ['src/services/api.js', 'trafficService, simulationService', 'REST API client for backend'],
            ['src/services/firebase.js', 'Firebase config', 'Firebase integration (optional data persistence)'],
        ],
        col_widths=[5, 5, 7]
    )

    add_styled_heading(doc, '4.3 Configuration & Infrastructure Files', 2)
    create_table(doc,
        ['File Path', 'Purpose'],
        [
            ['docker-compose.yml', 'Docker orchestration — defines backend (FastAPI, port 8000), camera server (Flask, port 5000), and frontend (React, port 3001) services'],
            ['sumo_files/', 'SUMO simulation network and route files for traffic simulation'],
            ['backend/app/services/ai_engine/model_ep*.pth', 'Saved PyTorch model weights from RL training episodes'],
        ],
        col_widths=[7, 11]
    )

    doc.add_page_break()

    # ========== 5. TECHNICAL SKILLS & TECHNOLOGIES ==========
    add_styled_heading(doc, '5. Technical Skills & Technologies Used', 1)

    create_table(doc,
        ['Category', 'Technology', 'Version / Details', 'Usage in UrbanFlow'],
        [
            ['AI / Machine Learning', 'PyTorch', 'Latest', 'Actor-Critic (A2C) neural network for RL agent — policy gradient training'],
            ['AI / Machine Learning', 'Reinforcement Learning', 'A2C / Actor-Critic', 'Learns optimal signal timing from traffic state observations and rewards'],
            ['AI / Machine Learning', 'Facebook Prophet', 'Latest', 'Time-series forecasting of future traffic volumes with confidence intervals'],
            ['Computer Vision', 'OpenCV (cv2)', '4.x', 'Frame processing, MOG2 background subtraction, Hough Circle Transform, morphology, contour detection, video encoding/decoding'],
            ['Computer Vision', 'YOLOv8 (Ultralytics)', 'Latest', 'Deep learning-based real-time vehicle detection and classification (car, motorcycle, bus, truck)'],
            ['Computer Vision', 'NumPy', 'Latest', 'Array operations for image processing, state vector computation, and normalization'],
            ['Backend Framework', 'FastAPI', 'Latest', 'Primary backend server — async REST API with Pydantic models, CORS middleware'],
            ['Backend Framework', 'Flask', 'Latest', 'Mobile camera server — serves camera HTML page, handles Socket.IO events for phone frames'],
            ['Backend Framework', 'Uvicorn', 'Latest', 'ASGI server for running FastAPI with async support'],
            ['Backend Framework', 'Eventlet', 'Latest', 'Async networking for Flask Socket.IO server'],
            ['Real-Time Communication', 'Socket.IO (Python)', 'Latest', 'Server-side WebSocket — async broadcast of traffic updates and AI decisions'],
            ['Real-Time Communication', 'Socket.IO (JavaScript)', '4.5.4', 'Client-side WebSocket — phone camera streaming and dashboard real-time updates'],
            ['Frontend Framework', 'React.js', '18.x', 'Component-based UI with hooks (useState, useEffect, useContext, useCallback)'],
            ['Frontend Framework', 'React Router', 'v6', 'Client-side routing — Home, Dashboard, Analytics pages'],
            ['UI Library', 'Material-UI (MUI)', 'v5/v6', 'Pre-built components — Paper, Typography, Grid, LinearProgress, ThemeProvider, CssBaseline'],
            ['Animation', 'Framer Motion', 'Latest', 'Smooth animations for ExplainabilityPanel — entrance animations, SHAP bar transitions'],
            ['Charting', 'Recharts', 'Latest', 'AreaChart (queue trends), BarChart (throughput) with gradients and responsive containers'],
            ['Mapping', 'React-Leaflet + Leaflet', 'Latest', 'Interactive map with traffic light markers, user geolocation, congestion circles, popups'],
            ['HTTP Client', 'Axios', 'Latest', 'REST API calls from frontend to FastAPI backend'],
            ['Data Validation', 'Pydantic', 'v2', 'Request/response model validation for FastAPI endpoints (TrafficData, SignalState, IntersectionStatus)'],
            ['Traffic Simulation', 'SUMO + TraCI', 'Latest', 'Microscopic traffic simulation — realistic vehicle generation, lane-level metrics, signal phase control via Python API'],
            ['Containerization', 'Docker + Docker Compose', 'Latest', 'Multi-service orchestration — backend, camera server, and frontend in isolated containers'],
            ['Data Science', 'Pandas', 'Latest', 'Data manipulation for Prophet training data preparation'],
            ['Cloud / BaaS', 'Firebase', 'Latest', 'Optional data persistence and authentication integration'],
            ['Language', 'Python 3.x', '3.10+', 'Backend logic, AI engine, computer vision, API server'],
            ['Language', 'JavaScript (ES6+)', 'ES2020+', 'Frontend React components, WebSocket client, camera streaming'],
            ['Language', 'HTML5 / CSS3', 'Latest', 'Mobile camera UI page, inline styles, responsive layout'],
        ],
        col_widths=[3, 3.5, 3, 8]
    )

    doc.add_page_break()

    # ========== 6. APIs ==========
    add_styled_heading(doc, '6. APIs — Internal & External', 1)

    add_styled_heading(doc, '6.1 REST APIs (FastAPI — Port 8000)', 2)
    create_table(doc,
        ['Method', 'Endpoint', 'Description', 'File'],
        [
            ['GET', '/', 'Health check — returns API status', 'app/main.py'],
            ['GET', '/health', 'Health endpoint — returns { status: "healthy" }', 'app/main.py'],
            ['GET', '/api/v1/traffic/', 'Get all intersection statuses', 'app/api/routes/traffic.py'],
            ['GET', '/api/v1/traffic/{id}', 'Get specific intersection status by ID', 'app/api/routes/traffic.py'],
            ['POST', '/api/v1/traffic/{id}/control', 'Update signal state (manual override or AI)', 'app/api/routes/traffic.py'],
            ['POST', '/api/v1/simulation/start', 'Start the simulation loop (simulation/real_world/hybrid mode)', 'app/api/routes/simulation.py'],
            ['POST', '/api/v1/simulation/stop', 'Stop the running simulation', 'app/api/routes/simulation.py'],
            ['GET', '/api/v1/simulation/status', 'Get current simulation status', 'app/api/routes/simulation.py'],
            ['POST', '/api/v1/simulation/camera-frame', 'Receive Base64 camera frame from phone/webcam', 'app/api/routes/simulation.py'],
            ['POST', '/api/v1/simulation/test-broadcast', 'Manually trigger a fake traffic update via WebSocket', 'app/api/routes/simulation.py'],
        ],
        col_widths=[2, 5, 6, 5]
    )

    add_styled_heading(doc, '6.2 WebSocket APIs (Socket.IO — Port 8000)', 2)
    create_table(doc,
        ['Event', 'Direction', 'Data', 'Description'],
        [
            ['connect', 'Client ← Server', '{ connected: true }', 'Fired when client connects. Starts broadcast loop on first connection.'],
            ['disconnect', 'Client ← Server', '—', 'Fired when client disconnects'],
            ['traffic_update', 'Server → Client', '{ intersections: [...], snapshot: { timestamp, avg_queue_length } }', 'Pushed every 2 seconds with full intersection state data'],
            ['ai_decision', 'Server → Client', '{ action_name, confidence, reasoning, feature_importance }', 'AI decision explanation for explainability panel'],
            ['emergency_alert', 'Server → Client', '{ type, message, ... }', 'Emergency vehicle priority notifications'],
            ['prediction_update', 'Server → Client', '{ predicted_vehicle_count, lower_bound, upper_bound }', 'Prophet forecast results'],
            ['metrics', 'Server → Client', '{ ... }', 'System-wide performance metrics'],
            ['connection_status', 'Bidirectional', '{ connected: bool }', 'Connection state synchronization'],
        ],
        col_widths=[3, 3, 6, 6]
    )

    add_styled_heading(doc, '6.3 Mobile Camera Server APIs (Flask — Port 5000)', 2)
    create_table(doc,
        ['Method / Event', 'Endpoint / Event Name', 'Description'],
        [
            ['GET', '/', 'Serves the camera HTML page to the phone browser'],
            ['GET', '/api/latest', 'Returns latest detection result (JSON)'],
            ['GET', '/api/stats', 'Returns camera and detector statistics'],
            ['GET', '/api/rl-state', 'Returns current 32-dimensional RL state vector'],
            ['GET', '/dashboard', 'Serves a monitoring dashboard page'],
            ['Socket.IO', 'frame (Client → Server)', 'Receives Base64-encoded video frame from phone'],
            ['Socket.IO', 'detection_result (Server → Client)', 'Sends detection results back to phone for overlay display'],
        ],
        col_widths=[3, 5, 10]
    )

    add_styled_heading(doc, '6.4 External Libraries & APIs Used', 2)
    create_table(doc,
        ['Library / API', 'What It Provides', 'How UrbanFlow Uses It'],
        [
            ['OpenCV (cv2)', 'Image processing, computer vision algorithms', 'Background subtraction (MOG2), contour detection, Hough circles, frame encode/decode, video capture/recording'],
            ['YOLOv8 (ultralytics)', 'Pre-trained object detection models', 'Detects vehicles with class labels (car=2, motorcycle=3, bus=5, truck=7) and confidence scores'],
            ['PyTorch', 'Neural network framework', 'ActorCritic model (nn.Module), loss computation, Adam optimizer, gradient backpropagation'],
            ['Facebook Prophet', 'Time-series forecasting', 'Predicts future vehicle counts from historical data with uncertainty intervals'],
            ['Leaflet + React-Leaflet', 'Interactive maps', 'MapContainer, TileLayer (OSM tiles), Marker, Circle, Popup for geospatial traffic visualization'],
            ['Recharts', 'React charting library', 'AreaChart, BarChart, ResponsiveContainer for real-time data visualization'],
            ['Framer Motion', 'Animation library', 'motion.div for entrance/exit animations, SHAP bar width transitions'],
            ['Material-UI (MUI)', 'React UI components', 'Paper, Grid, Typography, LinearProgress, ThemeProvider for consistent dark theme'],
            ['Axios', 'HTTP client', 'REST API calls: GET intersections, POST signal control, simulation start/stop'],
            ['Socket.IO', 'WebSocket abstraction', 'Bidirectional real-time communication between all system components'],
            ['Flask', 'Lightweight web framework', 'Camera server with HTML serving, REST API, and WebSocket support'],
            ['FastAPI', 'Modern Python web framework', 'Async REST API with automatic docs, Pydantic validation, CORS middleware'],
            ['SUMO TraCI', 'Traffic simulation control', 'Start/stop simulation, setPhase, getState, getReward via Python TraCI API'],
            ['Pydantic', 'Data validation', 'BaseModel classes for TrafficData, SignalState, IntersectionStatus'],
            ['Docker', 'Containerization', 'Multi-service deployment with docker-compose'],
        ],
        col_widths=[4, 5, 9]
    )

    doc.add_page_break()

    # ========== 7. DATA FLOW ==========
    add_styled_heading(doc, '7. Data Flow — End-to-End Cycle', 1)
    add_para(doc, 'One complete cycle of the UrbanFlow system (repeated every ~2 seconds):', bold=True)

    flow_steps = [
        ('t=0.0s  ', 'Phone captures video frame via browser camera access'),
        ('t=0.1s  ', 'JavaScript converts frame to Base64 JPEG and sends via Socket.IO to Camera Server (port 5000)'),
        ('t=0.2s  ', 'Camera Server decodes Base64 → OpenCV numpy array'),
        ('t=0.3s  ', 'VehicleDetector processes frame: "5 cars detected" | SignalDetector processes frame: "RED signal detected"'),
        ('t=0.4s  ', 'Lane assignment: N=2, S=1, E=1, W=1 | Speed estimation: 8.5 m/s avg | Queue: 0 stopped'),
        ('t=0.5s  ', 'Detection result emitted back to phone + /api/latest updated'),
        ('t=1.0s  ', 'Simulation step: RLAgent reads 32-dim state → Neural network outputs Action=0 (GREEN) → Action applied → Reward = -45 calculated'),
        ('t=2.0s  ', 'WebSocket broadcast: intersection data pushed to all connected dashboard clients'),
        ('t=2.0s  ', 'Dashboard updates: signal indicator changes, vehicle count updates, chart adds new data point, activity log entry added'),
        ('t=2.0s  ', 'Next phone frame arrives → cycle repeats ♻️'),
    ]
    for prefix, text in flow_steps:
        add_bullet(doc, text, bold_prefix=prefix)

    doc.add_page_break()

    # ========== 8. THREE OPERATING MODES ==========
    add_styled_heading(doc, '8. Three Operating Modes', 1)

    create_table(doc,
        ['Mode', 'Data Source', 'Use Case', 'Description'],
        [
            ['Simulation', 'SUMO Simulator (or Mock)', 'Testing & Training the AI', 'Generates synthetic traffic data. AI optimizes signal timing in a safe, repeatable environment. Dashboard shows simulated results.'],
            ['Real-World', 'Phone Camera → OpenCV', 'Live deployment at intersection', 'Real traffic video is processed by vehicle and signal detectors. AI suggests optimal timing. Dashboard shows live detection.'],
            ['Hybrid', '70% SUMO + 30% Camera', 'Calibrating simulation', 'Combines both sources with configurable weights. More robust than either alone. Used for transitioning from sim to real deployment.'],
        ],
        col_widths=[3, 4, 4, 7]
    )

    doc.add_page_break()

    # ========== 9. KEY INNOVATION ==========
    add_styled_heading(doc, '9. Key Innovation — Traditional vs. UrbanFlow', 1)

    create_table(doc,
        ['Aspect', 'Traditional Traffic Lights', 'UrbanFlow'],
        [
            ['Signal Timing', 'Fixed 30s green / 30s red always', 'Dynamic 15–60s based on real-time demand'],
            ['Traffic Awareness', 'No awareness of actual traffic', 'Sees actual traffic via camera + AI'],
            ['Adaptation', 'Same timing at 3 AM and 5 PM', 'Adapts in real-time — less wait at 3 AM, more green at 5 PM'],
            ['Data & Insights', 'No data, no visibility', 'Full analytics dashboard with historical trends'],
            ['Decision Making', 'Static rules, no learning', 'AI (A2C RL) that learns and improves over time'],
            ['Explainability', 'N/A', 'Natural language explanations + SHAP feature importance'],
            ['Scalability', 'Each intersection is independent', 'Multi-intersection coordination possible'],
            ['Cost', 'Dedicated hardware per intersection', 'Uses existing smartphone cameras'],
            ['Result', 'Wasted time and fuel', '30–40% less waiting time, reduced emissions'],
        ],
        col_widths=[3, 7, 7]
    )

    doc.add_page_break()

    # ========== 10. CONCLUSION ==========
    add_styled_heading(doc, '10. Conclusion', 1)
    add_para(doc,
        'UrbanFlow demonstrates a complete, end-to-end solution for intelligent traffic management. By combining '
        'accessible smartphone cameras, state-of-the-art computer vision, reinforcement learning, and a modern '
        'real-time web dashboard, it addresses all major pain points of traditional traffic signal systems:')

    conclusions = [
        'Eliminates blind, fixed-timer signals with AI-adaptive control',
        'Reduces waiting time by 30–40% through demand-responsive signal timing',
        'Provides complete traffic visibility via real-time dashboard and analytics',
        'Explains its decisions transparently with SHAP-like interpretability',
        'Scales affordably using existing phones instead of expensive hardware',
        'Supports gradual deployment from simulation to hybrid to real-world mode',
        'Reduces fuel waste and CO₂ emissions from idling vehicles',
    ]
    for c in conclusions:
        add_bullet(doc, c)

    add_para(doc, '')
    add_para(doc,
        'The project architecture features 14 backend modules and 13 frontend components, interconnected via '
        'REST APIs, WebSocket events, and a shared data pipeline. Built with Python (FastAPI, Flask, PyTorch, OpenCV) '
        'and JavaScript (React, Socket.IO, Recharts, Leaflet), UrbanFlow represents a production-ready prototype '
        'for the next generation of smart city traffic infrastructure.',
        italic=True)

    # Save
    output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'UrbanFlow_How_It_Solves_Traffic.docx')
    doc.save(output_path)
    print(f"\n✅ Document saved: {output_path}")
    return output_path


if __name__ == '__main__':
    generate_document()

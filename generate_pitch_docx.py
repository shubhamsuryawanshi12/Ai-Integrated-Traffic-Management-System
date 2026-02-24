
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

def set_cell_shading(cell, color_hex):
    shading = cell._element.get_or_add_tcPr()
    shade = shading.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear',
        qn('w:color'): 'auto',
        qn('w:fill'): color_hex
    })
    shading.append(shade)

def add_styled_heading(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        if level == 1:
            run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)
            run.font.size = Pt(24)
            run.bold = True
        elif level == 2:
            run.font.color.rgb = RGBColor(0x2C, 0x5F, 0x8A)
            run.font.size = Pt(18)
            run.bold = True
        elif level == 3:
            run.font.color.rgb = RGBColor(0x34, 0x7A, 0xB5)
            run.font.size = Pt(14)
            run.bold = True
    return heading

def add_para(doc, text, bold=False, italic=False, size=11, color=None, alignment=None, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = 'Calibri'
    if color:
        run.font.color.rgb = RGBColor(*color)
    if alignment:
        p.alignment = alignment
    p.paragraph_format.space_after = Pt(space_after)
    return p

def add_bullet(doc, text, bold_prefix=None, level=0):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.size = Pt(11)
        run = p.add_run(text)
        run.font.size = Pt(11)
    else:
        run = p.add_run(text)
        run.font.size = Pt(11)
    return p

def create_table(doc, headers, rows, col_widths=None, header_color='1E3A5F'):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, header_color)

    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r + 1].cells[c]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(10)
            if r % 2 == 0:
                set_cell_shading(cell, 'F0F4F8')

    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)

    doc.add_paragraph()
    return table

def generate_pitch_report():
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # TITLE PAGE
    for _ in range(4): doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('UrbanFlow: Pitching Report')
    run.font.size = Pt(36)
    run.bold = True
    run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('AI-Powered Adaptive Traffic Management System')
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x2C, 0x5F, 0x8A)

    for _ in range(6): doc.add_paragraph()
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run('Confidential Pitch Deck Supplement\nHackathon 2026')
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    doc.add_page_break()

    # 1. THE PROBLEM
    add_styled_heading(doc, '1. The Problem: Static Cities vs. Dynamic Traffic', 1)
    add_para(doc, "Urban congestion is a multi-billion dollar problem caused by a fundamental mismatch: traffic is dynamic, but control systems are static.")
    problems = [
        ("Economic Waste: ", "Billions lost annually in fuel and productivity."),
        ("Environmental Toll: ", "Vehicles idling at red lights contribute up to 20% of urban CO2 emissions."),
        ("Blind Systems: ", "Traditional lights operate on fixed timers, unaware of actual vehicle queues."),
        ("Frustration: ", "Drivers wait at empty intersections, leading to 'phantom' traffic jams.")
    ]
    for prefix, text in problems: add_bullet(doc, text, bold_prefix=prefix)

    # 2. THE SOLUTION
    add_styled_heading(doc, '2. The Solution: UrbanFlow', 1)
    add_para(doc, "UrbanFlow is an autonomous ecosystem that leverages Multi-Agent Reinforcement Learning (MARL) and Computer Vision to transform intersections into intelligent nodes.")
    add_para(doc, "Our system follows a cycle of PERCEPTION → INTELLIGENCE → ACTION.", bold=True)
    features = [
        ("Computer Vision: ", "Real-time vehicle detection using mobile/CCTV feeds with YOLOv8 & OpenCV."),
        ("AI Optimization: ", "MAPPO (Proximal Policy Optimization) agents that coordinate traffic phases."),
        ("Sustainability Focus: ", "Reduces global waiting time and CO2 emissions through 'Green Wave' optimization.")
    ]
    for prefix, text in features: add_bullet(doc, text, bold_prefix=prefix)

    # 3. TECHNICAL DEEP DIVE
    add_styled_heading(doc, '3. Technical Excellence', 1)
    add_styled_heading(doc, '3.1 AI Algorithm: MAPPO', 2)
    add_para(doc, "UrbanFlow uses Multi-Agent Proximal Policy Optimization. Unlike standard AI, MAPPO allows multiple intersections (agents) to cooperate.")
    create_table(doc, ["Component", "Technical Detail"], [
        ["Optimization", "PCU-based (Passenger Car Unit) Reward System"],
        ["Reward Function", "reward = -(waiting_time + 10 × stopped_vehicles)"],
        ["Neural Network", "Actor-Critic architecture with 128 hidden neurons (PyTorch)"],
        ["XAI (Explainable AI)", "SHAP-based feature importance for decision transparency"]
    ], col_widths=[6, 11])

    add_styled_heading(doc, '3.2 Perception Pipeline', 2)
    add_para(doc, "Our vision system is optimized for edge deployment using mobile cameras.")
    create_table(doc, ["Phase", "Methodology"], [
        ["Detection", "YOLOv8 + MOG2 Background Subtraction for robust tracking"],
        ["Classification", "Distinguishes between Cars, Trucks, Buses, and Motorcycles"],
        ["Signal Sense", "HSV Color Space analysis with Hough Circle Transform"],
        ["Metrics", "Calculates Queue Length, Avg Speed, and Throughput in real-time"]
    ], col_widths=[6, 11])

    # 4. DASHBOARD & OPERATIONAL VISIBILITY
    add_styled_heading(doc, '4. Dashboard & Features', 1)
    dash_features = [
        ("Real-time Analytics: ", "Live monitoring via React 19 and Leaflet Maps."),
        ("Explainability: ", "Typewriter-animated AI logic explaining 'Why' a signal was changed."),
        ("CO2 Tracker: ", "Visualizes fuel savings based on reduced idling intervals."),
        ("Emergency Mode: ", "Automated lane clearing for sirens/emergency-priority vehicles.")
    ]
    for prefix, text in dash_features: add_bullet(doc, text, bold_prefix=prefix)

    # 5. MARKET INNOVATION
    add_styled_heading(doc, '5. Traditional vs. UrbanFlow', 1)
    create_table(doc, ["Feature", "Traditional System", "UrbanFlow"], [
        ["Logic", "Fixed Timers", "AI-Adaptive"],
        ["Awareness", "Blind", "Real-world Perception"],
        ["Feedback Loop", "None", "Real-time Reward/Adaptation"],
        ["CO2 Impact", "High", "30-40% Reduction"],
        ["Hardware", "Expensive Sensors", "Standard Mobile/CCTV"]
    ], col_widths=[4, 6, 6])

    # 6. FUTURE VISION
    add_styled_heading(doc, '6. Scaling & Road-map', 1)
    roadmap = [
        "V2X Integration: Communicating directly with smart vehicles for zero-latency control.",
        "Edge Deployment: Running inference entirely on IoT chips at the pole.",
        "Cloud Coordination: City-wide synchronization via centralized Firebase hubs."
    ]
    for item in roadmap: add_bullet(doc, item)

    # CONCLUSION
    add_styled_heading(doc, '7. Conclusion', 1)
    add_para(doc, "UrbanFlow is not just a tool; it's the future infrastructure of smart cities. By combining multi-agent coordination with low-cost computer vision, we offer a scalable, sustainable solution to the world's traffic crisis.", italic=True)

    path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'UrbanFlow_Pitch_Report.docx')
    doc.save(path)
    return path

if __name__ == '__main__':
    path = generate_pitch_report()
    print(f"REPORT_GENERATED: {path}")
